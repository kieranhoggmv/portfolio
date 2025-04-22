import base64
import json
import pickle
from pprint import pprint
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import render
from django.views.generic import TemplateView
from docx import Document
import zipfile

KSBS = {
    "K1": {},
    "K2": {},
    "K5": {},
    "K6": {},
    "K7": {},
    "K10": {},
    "K13": {},
    "K14": {},
    "K15": {},
    "S5": {},
    "S9": {},
    "S10": {},
    "S11": {},
    "S13": {},
    "S14": {},
    "B1": {},
    "B2": {},
    "B5": {},
    "B6": {},
    "B7": {},
}


class Home(TemplateView):
    template_name = "upload.html"

    def post(self, request, *args, **kwargs):
        PROJECT_STYLE = "Heading 2"

        z = zipfile.ZipFile(request.FILES["file"].file)
        all_files = z.namelist()

        images = list(filter(lambda x: x.startswith("word/media/"), all_files))
        image_strings = []
        for image in images:
            filename = image.replace("word/media/", "")
            img = z.open(image).read()
            f = open(rf"images/{filename}", "wb")
            # f.write(img)
            image_strings.append(base64.b64encode(img))

            # z.extract(image, r"images")
        document = Document(request.FILES["file"].file)
        # paragraphs = [
        #     (para.text, para.style.name.replace(" ", "_"))
        #     for para in document.paragraphs
        # ]
        paragraphs = []
        projects = []
        project = None
        # previous_paragraph = None
        img = 1
        for i, paragraph in enumerate(document.paragraphs):
            if paragraph._p.xpath(
                "./w:r/w:drawing/*[self::wp:inline | self::wp:anchor]/a:graphic/a:graphicData/pic:pic"
            ):
                image = image_strings[img - 1].decode("utf-8")
                img += 1
            else:
                image = None
            for run in paragraph.runs:
                if run.font.italic:
                    print(run.text)
            if paragraph.style.name == PROJECT_STYLE and "#" in paragraph.text:
                if paragraph.text != project:
                    project = paragraph.text
                    projects.append(project)
            paragraphs.append(
                (paragraph.text, paragraph.style.name.replace(" ", "_"), image)
            )
            for key in KSBS:
                if (
                    f"[{key}]" in paragraph.text
                    or f"({key})" in paragraph.text
                    or f"/{key}" in paragraph.text
                    or f"{key}/" in paragraph.text
                    or f"{key}+" in paragraph.text
                    or f"{key}]" in paragraph.text
                ):
                    # if previous_paragraph:
                    #     KSBS[key]["paragraphs"].append(previous_paragraph.text)
                    if i not in KSBS[key].keys():
                        KSBS[key][i] = paragraph.text, project
            # previous_paragraph = paragraph
        # missing_ksbs = list(filter(lambda x: len(KSBS[x]) == 0, KSBS))
        # met_ksbs = list(filter(lambda x: x not in missing_ksbs, KSBS))
        # pprint(KSBS)
        criteria_file = open("data", "rb")
        criteria = pickle.load(criteria_file)
        for item in criteria:
            if item["ksb"] in KSBS:
                KSBS[item["ksb"]].update(item)
        print(KSBS)
        return render(
            request,
            self.template_name,
            {
                "ksbs": KSBS,
                # "missing_ksbs": missing_ksbs,
                # "met_ksbs": met_ksbs,
                "paragraphs": paragraphs,
                "criteria": criteria,
            },
        )
